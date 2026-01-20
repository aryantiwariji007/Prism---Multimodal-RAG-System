# Parse PDF 
from PyPDF2 import PdfReader
from docx import Document
import os
import sys
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# Add the backend directory to the Python path for progress service
backend_dir = Path(__file__).parent.parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

try:
    from app.services.progress_service import progress_service
except ImportError:
    # Fallback if progress service is not available
    progress_service = None

def parse_document(path, file_id=1, progress_file_id=None):
    """
    Parse either PDF or DOCX files based on file extension
    """
    file_extension = os.path.splitext(path)[1].lower()
    chunks = []
    
    if file_extension == '.pdf':
        chunks = parse_pdf(path, file_id, progress_file_id)
    elif file_extension == '.docx':
        chunks = parse_docx(path, file_id, progress_file_id)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}. Only PDF and DOCX files are supported.")
    
    return chunks

def parse_pdf(path, file_id=1, progress_file_id=None):
    """
    Parse PDF files using pdfplumber for "Enterprise Grade" extraction.
    Extracts proper tables and isolates them from text to prevent jumbled content.
    """

    chunks = []
    
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Opening PDF (Enterprise Mode)...")

    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            total_pages = len(pdf.pages)
            if progress_service and progress_file_id:
                progress_service.update_progress(progress_file_id, 1, f"Parsing {total_pages} pages...", 10)
            
            for i, page in enumerate(pdf.pages):
                # 1. Extract Tables with multiple strategies
                # Strategy A: Lines (for bordered tables)
                tables = page.extract_tables({
                    "vertical_strategy": "lines", 
                    "horizontal_strategy": "lines"
                })

                table_texts = []
                
                if tables:
                    for t_idx, table in enumerate(tables):
                        # Format:
                        # [Table]
                        # Row 1: Col1 | Col2
                        t_str = [f"\n--- Table {t_idx+1} (Page {i+1}) ---"]
                        for row in table:
                            # Filter None/Empty cells
                            cleaned_row = [str(cell).strip() for cell in row if cell and str(cell).strip()]
                            if cleaned_row:
                                t_str.append(" | ".join(cleaned_row))
                        t_str.append("--- End Table ---\n")
                        table_texts.append("\n".join(t_str))
                
                # 2. Extract Text with TIGHT tolerance to prevent column merging
                # x_tolerance=1 ensures "123  456" stays separate, but "H ell o" is merged.
                # Default is 3, which is often too loose for dense tables.
                raw_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                
                # Combine: Tables first (so they are prominent), then the rest of the text
                full_page_text = ""
                if table_texts:
                    full_page_text += "\n".join(table_texts) + "\n\n"
                
                full_page_text += raw_text
                
                chunks.append({
                    "file_id": file_id,
                    "page": i,
                    "text": full_page_text
                })

                # Progress
                if progress_service and progress_file_id:
                    page_progress = ((i + 1) / total_pages) * 90
                    progress_service.update_progress(
                        progress_file_id, 
                        1, 
                        f"Parsed page {i + 1}/{total_pages}", 
                        10 + page_progress
                    )
                    
    except (Exception, ImportError) as e:
        logger.error(f"pdfplumber failed: {e}. Falling back to PyPDF2.")
        # Fallback to PyPDF2 if pdfplumber fails (e.g. strict PDF)
        return _parse_pdf_fallback(path, file_id, progress_file_id)

    return chunks

def _parse_pdf_fallback(path, file_id, progress_file_id):
    """Original PyPDF2 implementation as fallback"""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(path)
        chunks = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception as e:
                logger.warning(f"Failed to extract text from page {i} of {path}: {e}")
                text = ""
            chunks.append({"file_id": file_id, "page": i, "text": text})
        return chunks
    except Exception as e:
        logger.error(f"Fallback PDF parsing failed for {path}: {e}")
        return [{"file_id": file_id, "page": 0, "text": ""}]

def parse_docx(path, file_id=1, progress_file_id=None):
    """Parse DOCX files using python-docx, including Tables"""
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Opening DOCX file...")
    
    doc = Document(path)
    chunks = []
    
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "Extracting text from document...", 20)
    
    # Extract text from all paragraphs
    full_text = []
    total_paragraphs = len(doc.paragraphs)
    
    # 1. Paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            full_text.append(paragraph.text)
        
        if progress_service and progress_file_id and total_paragraphs > 0:
            para_progress = ((i + 1) / total_paragraphs) * 40  # 40% of step 1 for paragraphs
            progress_service.update_progress(
                progress_file_id, 
                1, 
                f"Processing paragraph {i + 1} of {total_paragraphs}", 
                20 + para_progress
            )

    # 2. Tables (CRITICAL FIX for Questionnaires)
    total_tables = len(doc.tables)
    if total_tables > 0:
        logger.info(f"Found {total_tables} tables in DOCX")
        if progress_service and progress_file_id:
             progress_service.update_progress(progress_file_id, 1, f"Processing {total_tables} tables...", 60)

        for t_idx, table in enumerate(doc.tables):
            # Represent table as text
            # Table 1:
            # Row 1: Cell 1 | Cell 2
            table_text = [f"\n--- Table {t_idx + 1} ---"]
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    table_text.append(" | ".join(row_cells))
            
            if len(table_text) > 1: # If table has content
                full_text.extend(table_text)
                full_text.append("--- End Table ---\n")

    # Combine all text
    combined_text = '\n'.join(full_text)
    chunks.append({"file_id": file_id, "page": 0, "text": combined_text})
    
    if progress_service and progress_file_id:
        progress_service.update_progress(progress_file_id, 1, "DOCX parsing completed", 100)
    
    return chunks
