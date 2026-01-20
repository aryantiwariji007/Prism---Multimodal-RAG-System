import logging
import sys
from pathlib import Path
import pandas as pd
from docx import Document

# Setup path
backend_dir = Path(__file__).parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

from ingestion.parse_pdf import parse_docx
from ingestion.excel_ingestor import ingest_excel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("VERIFY")

def verify_docx_tables():
    print("\n--- Verifying DOCX Table Extraction ---")
    filename = "verify_table.docx"
    doc = Document()
    doc.add_paragraph("This is a normal paragraph.")
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Value A"
    table.cell(1, 1).text = "Section 6 Secret Content"
    
    doc.save(filename)
    
    chunks = parse_docx(filename, file_id="verify_docx")
    text = chunks[0]["text"]
    
    print(f"Extracted Text:\n{text[:500]}...")
    
    if "Section 6 Secret Content" in text:
        print("PASS: Table content found in DOCX chunk.")
    else:
        print("FAIL: 'Section 6 Secret Content' NOT found in DOCX chunk.")

    # Cleanup
    Path(filename).unlink()

def verify_excel_ingestion():
    print("\n--- Verifying Excel Ingestion ---")
    filename = "verify_sheet.xlsx"
    
    df = pd.DataFrame({
        "Product": ["Widget A", "Widget B"],
        "Price": [100, 200],
        "Notes": ["Expensive item", "Cheap item"]
    })
    
    df.to_excel(filename, index=False)
    
    chunks = ingest_excel(filename, file_id="verify_excel")
    text = chunks[0]["text"]
    
    print(f"Extracted Text:\n{text[:500]}...")
    
    # Check for row formatting
    if "Product: Widget A" in text and "Price: 100" in text:
        print("PASS: Row content found formatted correctly.")
    else:
        print("FAIL: Row content missing or malformed.")

    # Cleanup
    try:
        Path(filename).unlink()
    except:
        pass

if __name__ == "__main__":
    try:
        verify_docx_tables()
        verify_excel_ingestion()
    except Exception as e:
        print(f"Verification Failed: {e}")
